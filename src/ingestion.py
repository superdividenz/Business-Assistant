from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import os

def load_documents(file_paths):
    documents = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    
    for path in file_paths:
        if path.endswith('.pdf'):
            from PyPDF2 import PdfReader
            reader = PdfReader(path)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        elif path.endswith('.docx'):
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif path.endswith('.txt'):
            with open(path, 'r') as f:
                text = f.read()
        else:
            continue
        
        doc = Document(page_content=text, metadata={"source": path})
        docs = text_splitter.split_documents([doc])
        documents.extend(docs)
    
    return documents